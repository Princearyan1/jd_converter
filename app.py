from flask import Flask, render_template, request, send_file
import os
import fitz
from docx import Document
from groq import Groq
from dotenv import load_dotenv
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json

app = Flask(__name__)

load_dotenv()

client = Groq(
    api_key=os.getenv("GROQ_API_KEY")
)

def convert_jd_with_groq(text):

    prompt = f"""
You are a Deloitte Talent Acquisition Specialist.

Convert the Job Description into Deloitte format.
Extract and infer:
- Function
- Service Area
- Sub Service Area
- Primary Skill
- Grade/Designation

Return valid JSON only.



{{
  "function": "",
  "service_area": "",
  "sub_service_area": "",
  "skill": "",
  "grade": "",
  "team": "",
  "role_description": "",
  "responsibilities": [],
  "skills": [],
  "total_experience": "",
  "relevant_experience": "",
  "education": "",
  "certifications": [],
  "location": "",
  "travel_requirement": ""
}}


Generate minimum 8 responsibilities.
Generate minimum 10 skills.
Return valid JSON only.
Job Description:

{text}
"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        temperature=0.2,
        messages=[
            {
                "role":"user",
                "content":prompt
            }
        ]
    )

    result = response.choices[0].message.content.strip()

    if result.startswith("```json"):
        result = result.replace("```json","")
        result = result.replace("```","")

    return json.loads(result)
def create_deloitte_doc(data):

    doc = Document()
    section = doc.sections[0]

    hdr = section.header
    
    # Deloitte Left Side
    p1 = hdr.paragraphs[0]
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    r1 = p1.add_run("Deloitte")
    r1.bold = True
    r1.font.size = Pt(20)
    r1.font.color.rgb = RGBColor(1,2,3)
    r3 = p1.add_run(".")
    r3.bold = True
    r3.font.size = Pt(20)
    r3.font.color.rgb = RGBColor(134,188,37)
    
    # New Paragraph for Heading
    p2 = hdr.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    r2 = p2.add_run(
        f"{data.get('function','Risk Advisory')} | "
        f"{data.get('service_area','Cyber')} | "
        f"{data.get('sub_service_area','Risk')} | "
        f"{data.get('skill','Technology')} | "
        f"{data.get('grade','Assistant Manager')}"
    )
    
    r2.bold = True
    r2.font.size = Pt(13)


    


    doc.add_paragraph()
    doc.add_heading("Your potential, unleashed.", level=1)

    

    doc.add_paragraph(
        "India’s impact on the global economy has increased at an exponential rate and Deloitte presents an opportunity to unleash and realize your potential amongst cutting edge leaders, and organizations shaping the future of the region, and indeed, the world beyond."
    )

    doc.add_paragraph(
        "At Deloitte, bring your whole self to work, every day. Combine that with our drive to propel with purpose and you have the perfect playground to collaborate, innovate, grow, and make an impact that matters."
    )

    doc.add_heading("The Team", level=1)

    doc.add_paragraph(
        f"{data.get('team','Risk Advisory')} Team is about much more than just the numbers. "
        "It’s about attesting to accomplishments and challenges and helping to assure strong foundations for future aspirations. Deloitte exemplifies what, how, and why of change so you’re always ready to act ahead."
    )

    doc.add_heading("Your Work Profile", level=1)

    doc.add_paragraph(
        f"As an {data.get('grade','Assistant Manager')} in our "
        f"{data.get('team','Risk Advisory')} team you’ll build and nurture positive working relationships with teams and clients with the intention to exceed client expectations."
    )

    doc.add_paragraph(
        "As a part of our Risk Advisory team, you’ll build and nurture positive working relationships with teams and clients with the intention to exceed client expectations."
    )

    doc.add_paragraph(
        f"Role Description - {data.get('role_description','')}"
    )

    for item in data.get("responsibilities", []):
        doc.add_paragraph(
            item,
            style="List Bullet"
        )

    doc.add_paragraph(
        "As a prospective candidate, you should possess:"
    )

    for skill in data.get("skills", []):
            doc.add_paragraph(
                skill,
                style="List Bullet"
            )
    
    doc.add_heading(
            "Desired Qualifications",
            level=1
        )
    
    doc.add_paragraph(
            f"Total experience – {data.get('total_experience','')}"
        )
    
    doc.add_paragraph(
            f"Relevant experience – {data.get('relevant_experience','')}"
        )
    
    doc.add_paragraph(
            f"Education – {data.get('education','')}"
        )
    
    for cert in data.get("certifications", []):
            doc.add_paragraph(
                cert,
                style="List Bullet"
            )
    
    doc.add_heading(
            "Location and way of working",
            level=1
        )
    
    doc.add_paragraph(
            f"Base location: {data.get('location','PAN India')}"
        )
    
    doc.add_paragraph(
            data.get(
                'travel_requirement',
                'This profile does not involve extensive travel for work.'
            )
        )
    
    doc.add_paragraph(
            "Hybrid is our default way of working. Each domain has customized the hybrid approach to their unique needs."
        )
    
    doc.add_heading(
            f"Your role as {data.get('grade','Assistant Manager')}",
            level=1
        )
    
    traits = [
            "Inspiring - Leading with integrity to build inclusion and motivation.",
            "Committed to creating purpose - Creating a sense of vision and purpose.",
            "Agile - Achieving high-quality results through collaboration and Team unity.",
            "Skilled at building diverse capability - Developing diverse capabilities for the future.",
            "Persuasive / Influencing - Persuading and influencing stakeholders.",
            "Collaborating - Partnering to build new solutions.",
            "Delivering value - Showing commercial acumen.",
            "Committed to expanding business - Leveraging new business opportunities.",
            "Analytical Acumen - Leveraging data to recommend impactful approach and solutions.",
            "Effective communication – Must be well abled to have well-structured and well-articulated conversations.",
            "Engagement Management / Delivery Excellence - Effectively managing engagements.",
            "Managing change - Responding to changing environment with resilience.",
            "Managing Quality & Risk - Delivering high quality results and mitigating risks.",
            "Strategic Thinking & Problem Solving - Applying strategic mindset to solve business issues.",
            "Tech Savvy - Leveraging ethical technology practices.",
            "Empathetic leadership and inclusivity - Creating a safe and thriving environment."
        ]
    
    for t in traits:
            doc.add_paragraph(
                t,
                style="List Bullet"
            )
    
    output_file = os.path.join(
            OUTPUT_FOLDER,
            "Deloitte_JD.docx"
        )
    
    doc.save(output_file)
    
    
    return output_file





UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["OUTPUT_FOLDER"] = OUTPUT_FOLDER


def extract_pdf_text(path):

    doc = fitz.open(path)

    text = ""

    for page in doc:
        text += page.get_text()

    return text


def extract_docx_text(path):

    doc = Document(path)

    text = ""

    for para in doc.paragraphs:
        text += para.text + "\n"

    return text


@app.route("/")
def home():

    return render_template("index.html")



@app.route("/generate", methods=["POST"])
def generate():

    if "jd_file" not in request.files:
        return "No file selected"

    file = request.files["jd_file"]

    if file.filename == "":
        return "No file selected"

    filepath = os.path.join(
        app.config["UPLOAD_FOLDER"],
        file.filename
    )

    file.save(filepath)

    ext = file.filename.split(".")[-1].lower()

    if ext == "pdf":

        extracted_text = extract_pdf_text(filepath)

    elif ext == "docx":

        extracted_text = extract_docx_text(filepath)

    else:

        return "Only PDF and DOCX allowed"

    try:

        data = convert_jd_with_groq(
            extracted_text[:15000]
        )

        doc_path = create_deloitte_doc(data)

        return send_file(
            doc_path,
            as_attachment=True,
            download_name="Deloitte_JD.docx"
        )

    except Exception as e:

        return f"""
        <h2>Error</h2>
        <pre>{str(e)}</pre>
        """


if __name__ == "__main__":

    app.run(
        debug=True
    )